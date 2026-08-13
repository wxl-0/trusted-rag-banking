import pytest

from src.parser.base import Chunk


def test_chunk_preserves_cell_metadata():
    chunk = Chunk(
        doc_id="STAT-001",
        chunk_id="STAT-001#Sheet1#C5",
        text="2025-09, Sheet1 C5: row=Premium, column=YTD, value=123.45 yuan.",
        chunk_type="table_row",
        source_title="Monthly report",
        issuer="NFRA",
        doc_no="",
        publish_date="2025-09-01",
        section_path=[],
        source_url="",
        local_path="data/raw/report.xlsx",
        table_name="Sheet1",
        indicator="Premium",
        period="2025-09",
        unit="unit: yuan",
        row_index=5,
        cell_ref="C5",
        row_label="Premium",
        column_header="YTD",
        raw_value="123.45",
    )

    data = chunk.to_dict()

    assert data["cell_ref"] == "C5"
    assert data["row_label"] == "Premium"
    assert data["column_header"] == "YTD"
    assert data["raw_value"] == "123.45"


def test_chunk_from_dict_ignores_unknown_fields():
    data = {
        "doc_id": "TEST-001",
        "chunk_id": "TEST-001#body",
        "text": "body text",
        "chunk_type": "clause",
        "source_title": "Test",
        "issuer": "NFRA",
        "doc_no": "",
        "publish_date": "2024-01-01",
        "section_path": [],
        "source_url": "",
        "local_path": "",
        "unknown_payload_key": "kept out",
    }

    restored = Chunk.from_dict(data)

    assert restored.doc_id == "TEST-001"
    assert not hasattr(restored, "unknown_payload_key")


def test_excel_parser_returns_cell_level_chunks(tmp_path):
    import openpyxl
    from src.parser.excel_parser import ExcelParser

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Monthly"
    sheet.append(["", "2025 September insurance report", ""])
    sheet.append(["", "", "unit: yuan"])
    sheet.append(["", "metric", "YTD"])
    sheet.append(["", "Premium income", 123.45])
    path = tmp_path / "report.xlsx"
    workbook.save(path)

    chunks = ExcelParser(
        doc_id="STAT-001",
        source_title="2025 September insurance report",
        issuer="NFRA",
        publish_date="2025-09-01",
        source_url="",
        local_path=str(path),
    ).parse()

    assert len(chunks) == 1
    assert chunks[0].chunk_type == "table_row"
    assert chunks[0].cell_ref == "C4"
    assert chunks[0].row_label == "Premium income"
    assert chunks[0].column_header == "YTD"
    assert chunks[0].raw_value == "123.45"
    assert chunks[0].unit == "unit: yuan"


def test_excel_parser_preserves_grouped_row_labels(tmp_path):
    import openpyxl
    from src.parser.excel_parser import ExcelParser

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "机构分类"
    sheet.append(["机构", None, "大型商业银行", "外资银行"])
    sheet.append(["时间/指标", None, None, None])
    sheet.append(["一季度", "不良贷款余额", 12461.06, 121.77])
    sheet.append([None, "次级类贷款余额", 6032.98, 55.4])
    path = tmp_path / "grouped-rows.xlsx"
    workbook.save(path)

    chunks = ExcelParser(
        doc_id="STAT-GROUPED",
        source_title="2023年商业银行主要指标分机构类情况表",
        issuer="NFRA",
        publish_date="2023-12-31",
        source_url="",
        local_path=str(path),
    ).parse()
    by_cell = {chunk.cell_ref: chunk for chunk in chunks}

    assert by_cell["C3"].row_label == "不良贷款余额"
    assert by_cell["C3"].column_header == "大型商业银行"
    assert by_cell["C3"].section_path == ["一季度"]
    assert by_cell["C4"].row_label == "次级类贷款余额"
    assert by_cell["C4"].section_path == ["一季度"]


def test_excel_parser_preserves_quarters_and_repeated_table_sections(tmp_path):
    import openpyxl
    from src.parser.excel_parser import ExcelParser

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "资产负债季度"
    sheet.append(["银行业金融机构资产负债情况表(季度)", None, None])
    sheet.append(["1. 银行业金融机构", None, None])
    sheet.append(["时间", "2023年", None])
    sheet.append(["项目", "一季度", "四季度"])
    sheet.append(["总负债", 3648440.27, 3831244.74])
    sheet.append([None, None, None])
    sheet.append(["其中：商业银行合计", None, None])
    sheet.append(["时间", "2023年", None])
    sheet.append(["项目", "一季度", "四季度"])
    sheet.append(["总资产", 3371160.25, 3548466.70])
    path = tmp_path / "repeated-sections.xlsx"
    workbook.save(path)

    chunks = ExcelParser(
        doc_id="STAT-SECTIONS",
        source_title="2023年银行业总资产、总负债（季度）",
        issuer="NFRA",
        publish_date="2023-12-31",
        source_url="",
        local_path=str(path),
    ).parse()
    by_cell = {chunk.cell_ref: chunk for chunk in chunks}

    assert by_cell["B5"].column_header == "一季度"
    assert by_cell["C5"].column_header == "四季度"
    assert by_cell["B5"].section_path == ["1. 银行业金融机构"]
    assert by_cell["B10"].row_label == "总资产"
    assert by_cell["C10"].section_path == ["其中：商业银行合计"]


def test_excel_parser_supports_xls_files(tmp_path):
    pytest.importorskip("xlrd")
    xlwt = pytest.importorskip("xlwt")
    from src.parser.excel_parser import ExcelParser

    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("Legacy")
    rows = [
        ["", "2024 legacy report", ""],
        ["", "", "unit: yuan"],
        ["", "metric", "YTD"],
        ["", "Premium income", 88.5],
    ]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            sheet.write(row_idx, col_idx, value)
    path = tmp_path / "legacy.xls"
    workbook.save(str(path))

    chunks = ExcelParser(
        doc_id="STAT-002",
        source_title="2024 legacy report",
        issuer="NFRA",
        publish_date="2024-01-01",
        source_url="",
        local_path=str(path),
    ).parse()

    assert len(chunks) == 1
    assert chunks[0].cell_ref == "C4"
    assert chunks[0].row_label == "Premium income"
    assert chunks[0].column_header == "YTD"
    assert chunks[0].raw_value == "88.5"


def test_pdf_parser_keeps_heading_and_numbered_item_text_across_pages(tmp_path):
    import fitz
    from src.parser.pdf_parser import PdfParser

    path = tmp_path / "group-list.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text(
        (50, 60), "一、保险控股型集团", fontsize=16, fontname="china-s"
    )
    page.insert_text(
        (50, 100), "下列集团应当按照监管规则", fontsize=18, fontname="china-s"
    )
    page.insert_text(
        (50, 120), "编报报告。", fontsize=18, fontname="china-s"
    )
    page.insert_text(
        (50, 160), "（一）中国人民保险集团股份有限公司", fontsize=18,
        fontname="china-s",
    )
    page = document.new_page()
    page.insert_text(
        (50, 60), "或其指定的成员公司。", fontsize=12, fontname="china-s"
    )
    company_names = [
        "（二）中国人寿保险（集团）公司",
        "（三）中国太平保险集团有限责任公司",
        "（五）中国平安保险（集团）股份有限公司",
    ]
    for index, company_name in enumerate(company_names, start=1):
        page.insert_text(
            (50, 60 + index * 40), company_name,
            fontsize=18, fontname="china-s",
        )
    document.save(path)
    document.close()

    chunks = PdfParser(
        doc_id="PDF-LIST",
        source_title="应当编报保险集团偿付能力报告的公司名单",
        issuer="NFRA",
        doc_no="",
        publish_date="2021-12-30",
        source_url="",
        local_path=str(path),
    ).parse()

    assert any("一、保险控股型集团" in chunk.text for chunk in chunks)
    assert any(
        "下列集团应当按照监管规则 编报报告" in chunk.text
        and chunk.section_path == ["一、保险控股型集团"]
        for chunk in chunks
    )
    assert any(
        "（一）中国人民保险集团股份有限公司" in chunk.text
        and "或其指定的成员公司" in chunk.text
        and chunk.section_path == [
            "一、保险控股型集团",
            "（一）中国人民保险集团股份有限公司",
        ]
        for chunk in chunks
    )
    for company_name in company_names:
        assert any(
            company_name in chunk.text
            and chunk.section_path == ["一、保险控股型集团", company_name]
            for chunk in chunks
        )


def test_word_parser_includes_tables(tmp_path):
    from docx import Document
    from src.parser.word_parser import WordParser

    document = Document()
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("Paragraph evidence.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Rule"
    table.cell(1, 0).text = "Retention"
    table.cell(1, 1).text = "Keep for 10 years"
    path = tmp_path / "test.docx"
    document.save(path)

    chunks = WordParser(
        doc_id="DOC-001",
        source_title="Rules",
        issuer="NFRA",
        doc_no="",
        publish_date="2024-01-01",
        source_url="",
        local_path=str(path),
    ).parse()

    assert any(c.chunk_type == "clause" and "Paragraph evidence" in c.text for c in chunks)
    table_chunks = [c for c in chunks if c.chunk_type == "table_row"]
    assert len(table_chunks) == 1
    assert table_chunks[0].row_label == "Retention"
    assert table_chunks[0].column_header == "Item | Rule"
    assert "Keep for 10 years" in table_chunks[0].text


def test_word_parser_rejects_doc_without_converter(tmp_path, monkeypatch):
    from src.parser.word_parser import WordParser

    path = tmp_path / "legacy.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0")
    parser = WordParser(
        doc_id="DOC-002",
        source_title="Legacy",
        issuer="NFRA",
        doc_no="",
        publish_date="2024-01-01",
        source_url="",
        local_path=str(path),
    )
    monkeypatch.setattr(parser, "_convert_doc_to_docx", lambda _: None)

    with pytest.raises(RuntimeError, match="Unable to convert .doc"):
        parser.parse()


def test_pdf_parser_returns_page_level_paragraphs(tmp_path):
    import fitz
    from src.parser.pdf_parser import PdfParser

    path = tmp_path / "test.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((50, 50), "Chapter 1 General Rules", fontsize=16)
    page.insert_text((50, 95), "Article 1 This paragraph is evidence.", fontsize=12)
    document.save(path)
    document.close()

    chunks = PdfParser(
        doc_id="PDF-001",
        source_title="PDF Rules",
        issuer="NFRA",
        doc_no="",
        publish_date="2024-01-01",
        source_url="",
        local_path=str(path),
    ).parse()

    assert len(chunks) == 1
    assert chunks[0].page_no == 1
    assert chunks[0].section_path == ["Chapter 1 General Rules"]
    assert "Article 1" in chunks[0].text
