import pytest
from src.parser.base import Chunk


def test_chunk_to_dict_clause():
    chunk = Chunk(
        doc_id="NFRA-001",
        chunk_id="NFRA-001#第三章#第十二条",
        text="商业银行资本充足率不得低于10.5%。",
        chunk_type="clause",
        source_title="商业银行资本管理办法",
        issuer="国家金融监督管理总局",
        doc_no="银监发〔2023〕1号",
        publish_date="2023-11-01",
        section_path=["第三章 资本充足率", "第十二条"],
        source_url="https://www.nfra.gov.cn/xxx",
        local_path="data/raw/NFRA-001.pdf",
    )
    d = chunk.to_dict()
    assert d["doc_id"] == "NFRA-001"
    assert d["chunk_type"] == "clause"
    assert "table_name" not in d


def test_chunk_to_dict_table_row():
    chunk = Chunk(
        doc_id="STAT-001",
        chunk_id="STAT-001#Sheet1#R5",
        text="2024年三季度末，不良贷款率为1.56%。",
        chunk_type="table_row",
        source_title="G11资产质量情况表",
        issuer="国家金融监督管理总局",
        doc_no="",
        publish_date="2024-09-30",
        section_path=[],
        source_url="https://www.nfra.gov.cn/yyy",
        local_path="data/raw/STAT-001.xlsx",
        table_name="G11《资产质量情况表》",
        indicator="不良贷款率",
        period="2024Q3",
        unit="%",
        row_index=5,
    )
    d = chunk.to_dict()
    assert d["chunk_type"] == "table_row"
    assert d["indicator"] == "不良贷款率"
    assert d["period"] == "2024Q3"


def test_chunk_from_dict_roundtrip():
    chunk = Chunk(
        doc_id="TEST-001",
        chunk_id="TEST-001#body",
        text="测试文本",
        chunk_type="clause",
        source_title="测试文件",
        issuer="测试机构",
        doc_no="",
        publish_date="2024-01-01",
        section_path=[],
        source_url="",
        local_path="",
    )
    d = chunk.to_dict()
    restored = Chunk.from_dict(d)
    assert restored.doc_id == chunk.doc_id
    assert restored.text == chunk.text


def test_word_parser_returns_chunks(tmp_path):
    from docx import Document
    from src.parser.word_parser import WordParser

    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_heading("第一条", level=2)
    doc.add_paragraph("本办法适用于在中华人民共和国境内依法设立的商业银行。")
    test_file = tmp_path / "test.docx"
    doc.save(str(test_file))

    parser = WordParser(
        doc_id="TEST-001",
        source_title="测试办法",
        issuer="测试机构",
        doc_no="测试〔2024〕1号",
        publish_date="2024-01-01",
        source_url="https://example.com",
        local_path=str(test_file),
    )
    chunks = parser.parse()
    assert len(chunks) >= 1
    assert all(c.chunk_type == "clause" for c in chunks)
    assert all(c.doc_id == "TEST-001" for c in chunks)
    assert any("商业银行" in c.text for c in chunks)


def test_pdf_parser_returns_chunks(tmp_path):
    import fitz
    from src.parser.pdf_parser import PdfParser

    pdf_path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "第一章 总则", fontsize=16)
    page.insert_text((50, 100), "第一条 本办法适用于商业银行。", fontsize=12)
    doc.save(str(pdf_path))
    doc.close()

    parser = PdfParser(
        doc_id="TEST-002",
        source_title="测试PDF",
        issuer="测试机构",
        doc_no="",
        publish_date="2024-01-01",
        source_url="https://example.com",
        local_path=str(pdf_path),
    )
    chunks = parser.parse()
    assert len(chunks) >= 1
    assert all(c.chunk_type == "clause" for c in chunks)


def test_excel_parser_returns_table_chunks(tmp_path):
    import openpyxl
    from src.parser.excel_parser import ExcelParser

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "G11资产质量"
    ws.append(["指标名称", "2024Q3", "单位"])
    ws.append(["不良贷款率", 1.56, "%"])
    ws.append(["关注类贷款率", 2.34, "%"])
    xlsx_path = tmp_path / "test.xlsx"
    wb.save(str(xlsx_path))

    parser = ExcelParser(
        doc_id="STAT-001",
        source_title="G11资产质量情况表",
        issuer="国家金融监督管理总局",
        publish_date="2024-09-30",
        source_url="https://example.com",
        local_path=str(xlsx_path),
    )
    chunks = parser.parse()
    assert len(chunks) == 2
    assert all(c.chunk_type == "table_row" for c in chunks)
    assert chunks[0].indicator == "不良贷款率"
    assert "1.56" in chunks[0].text
